from docx import Document
from docxcompose.composer import Composer
import csv
import glob, os
from docx.shared import Pt

files_list = []

def csvReader():
    with open('data.csv') as csv_file:
        csv_reader = csv.reader(csv_file, delimiter=',')
        line_count = 0
        for row in csv_reader:
            if line_count == 0:
                line_count += 1
            else:
                print(f'\t{row[0]} {row[1]} email address is {row[2]} and the password is, {row[3]}.')
                line_count += 1
                document = Document('template.docx')
                count = 0
                for para in document.paragraphs:
                    if 'First name' in para.text:
                        print(para.text + f"\t{row[0]}")
                        para.add_run(f"\t{row[0]}")
                    elif 'Last name' in para.text:
                        print(para.text + f"\t{row[1]}")
                        para.add_run(f"\t{row[1]}")
                    elif 'Class' in para.text and count > 6:
                        print(para.text + f"\t{row[4]}")
                        para.add_run(f"\t{row[4]}")
                    elif 'Username' in para.text:
                        print(para.text + f"\t{row[2]}")
                        para.add_run(f"\t{row[2]}")
                    elif 'Password' in para.text:
                        print(para.text + f"\t{row[3]}")
                        para.add_run(f"\t{row[3]}")
                    count = count + 1
                document.save(f"{row[0]}_{row[1]}_GoogleInfo.docx")
        
        print(f'Processed {line_count} lines.')
        files_In_Directory()
    
def files_In_Directory():
    os.chdir("./")
    for file in glob.glob("*.docx"):
        if file != 'template.docx':
            files_list.append(file)
    combine_all_docx(files_list)


    
def combine_all_docx(files_list):
    print("Please Wait... Combining")
    number_of_sections=len(files_list)
    master = Document("template.docx")
    composer = Composer(master)
    for i in range(0, number_of_sections):
        doc_temp = Document(files_list[i])
        composer.append(doc_temp)
    composer.save("combined_file.docx")
    print("Done!")
    cleanup(files_list)

def cleanup(files_list):
    print("Please Wait... Cleaning Up")

    try: 
        pos = files_list.index("combined_file.docx")
        files_list.remove(pos)
    except:
        pass


    i = 0
    while i < len(files_list):
        print("Removing: "+files_list[i])
        os.remove(files_list[i])
        i = i + 1

    print(" DONE ")

csvReader()